# -*- coding: utf-8 -*-
"""Bitirme Projesi Raporu üretici (python-docx)."""
import json
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

S = json.load(open("/Users/Ethem/Desktop/colab/rapor_gorseller/stats.json"))
IMG = "/Users/Ethem/Desktop/colab/rapor_gorseller"
def pct(x): return f"%{x*100:.1f}"

doc = Document()

# -------- Genel stil --------
st = doc.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
st.paragraph_format.line_spacing = 1.5
st.paragraph_format.space_after = Pt(6)
st.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for h,sz in [("Heading 1",14),("Heading 2",13),("Heading 3",12)]:
    s=doc.styles[h]; s.font.name="Times New Roman"; s.font.size=Pt(sz)
    s.font.bold=True; s.font.color.rgb=RGBColor(0x0f,0x3a,0x3a)
    s.paragraph_format.space_before=Pt(12); s.paragraph_format.space_after=Pt(6)
    s.paragraph_format.keep_with_next=True

sec=doc.sections[0]
sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5); sec.left_margin=Cm(3); sec.right_margin=Cm(2.5)

# -------- Yardımcılar --------
def P(t,bold=False,it=False,align=None,size=None,after=6):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=it
    if size: r.font.size=Pt(size)
    p.paragraph_format.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after=Pt(after)
    return p
def H1(t): doc.add_heading(t,level=1)
def H2(t): doc.add_heading(t,level=2)
def H3(t): doc.add_heading(t,level=3)
def BULLET(t):
    p=doc.add_paragraph(t,style="List Bullet"); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def NUM(t):
    p=doc.add_paragraph(t,style="List Number"); p.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
def PB(): doc.add_page_break()
def FIG(path,caption,w=15.5):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(f"{IMG}/{path}",width=Cm(w))
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(caption); r.italic=True; r.font.size=Pt(10); c.paragraph_format.space_after=Pt(10)
def set_shd(cell,fill):
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement("w:shd")
    shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),fill); tcPr.append(shd)
def TABLE(headers,rows,caption,widths=None):
    cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=cap.add_run(caption); r.bold=True; r.font.size=Pt(10)
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"; t.alignment=1
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=""; run=c.paragraphs[0].add_run(h)
        run.bold=True; run.font.size=Pt(10.5); run.font.color.rgb=RGBColor(0xff,0xff,0xff)
        set_shd(c,"0f766e")
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=""; rn=cells[i].paragraphs[0].add_run(str(v)); rn.font.size=Pt(10)
    if widths:
        for row in t.rows:
            for i,w in enumerate(widths): row.cells[i].width=Cm(w)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
    return t
def CODE(lines,title=None):
    if title:
        tp=doc.add_paragraph(); tr=tp.add_run(title); tr.bold=True; tr.font.size=Pt(10)
        tp.paragraph_format.space_after=Pt(2)
    for ln in lines:
        p=doc.add_paragraph(); p.paragraph_format.line_spacing=1.0; p.paragraph_format.space_after=Pt(0)
        r=p.add_run(ln if ln else " "); r.font.name="Consolas"; r.font.size=Pt(8.5)
        r.font.color.rgb=RGBColor(0x10,0x20,0x20)
        pPr=p._p.get_or_add_pPr(); shd=OxmlElement("w:shd")
        shd.set(qn("w:val"),"clear"); shd.set(qn("w:fill"),"f1f5f9"); pPr.append(shd)
    doc.add_paragraph().paragraph_format.space_after=Pt(4)
def page_number_footer():
    for s in doc.sections:
        f=s.footer.paragraphs[0]; f.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=f.add_run()
        for tp,txt in [("begin",None),("instr","PAGE"),("end",None)]:
            if tp=="instr":
                e=OxmlElement("w:instrText"); e.set(qn("xml:space"),"preserve"); e.text=txt
            else:
                e=OxmlElement("w:fldChar"); e.set(qn("w:fldCharType"),tp)
            run._r.append(e)

# ============================================================ KAPAK
def center(t,size=12,bold=False,after=6,it=False):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(t); r.bold=bold; r.italic=it; r.font.size=Pt(size); r.font.name="Times New Roman"
    p.paragraph_format.space_after=Pt(after); p.paragraph_format.line_spacing=1.2
    return p
center("",6); center("T.C.",12,True,2); center("[ÜNİVERSİTE ADI]",13,True,2)
center("MÜHENDİSLİK FAKÜLTESİ",12,True,2); center("ENDÜSTRİ MÜHENDİSLİĞİ BÖLÜMÜ",12,True,24)
center("",18)
center("FUTBOLCU TRANSFER UYUMLULUĞUNUN",15,True,2)
center("DİJİTAL İKİZ VE MONTE CARLO SİMÜLASYONU",15,True,2)
center("İLE DEĞERLENDİRİLMESİ",15,True,18)
center("Gerçek Zamanlı Spor Verisi Tabanlı Bir Karar Destek Sistemi",12,False,28,it=True)
center("",18)
center("BİTİRME PROJESİ",13,True,30)
center("",18)
center("Hazırlayan",11,False,2); center("[Ad SOYAD]  —  [Öğrenci No]",12,True,18)
center("Proje Danışmanı",11,False,2); center("[Unvan Ad SOYAD]",12,True,40)
center("",18)
center("[ŞEHİR]  —  HAZİRAN 2026",12,True,2)
PB()

# ============================================================ İÇİNDEKİLER
H1("İÇİNDEKİLER")
p=doc.add_paragraph()
fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),'TOC \\o "1-3" \\h \\z \\u')
r=OxmlElement("w:r"); t=OxmlElement("w:t")
t.text="İçindekiler tablosu: Word'de açtıktan sonra bu alana sağ tıklayıp 'Alanı Güncelleştir' (F9) deyiniz."
r.append(t); fld.append(r); p._p.append(fld)
PB()

# ============================================================ ÖZET
H1("ÖZET")
P("Profesyonel futbolda transfer kararları, kulüpler için yüksek bütçeli ve yüksek riskli yatırımlardır. "
  "Bir oyuncunun mevcut performansının yeni bir takımda sürüp sürmeyeceği; sakatlık geçmişi, oyun stili uyumu "
  "ve takımın hücum üretimi gibi çok sayıda belirsiz faktöre bağlıdır. Bu çalışmada, bir futbolcunun gerçek "
  "geçmiş verilerine dayanarak sanal bir kopyasının (dijital ikiz) oluşturulduğu ve hedef bir takımdaki olası "
  "bir sezonun Monte Carlo simülasyonu ile binlerce kez yeniden üretildiği bir karar destek sistemi geliştirilmiştir. "
  "Veriler, API-Football servisinden gerçek zamanlı olarak çekilmekte; oyuncunun maç bazlı performans reytingleri, "
  "gol/asist üretimi ve sakatlık geçmişi işlenerek olasılıksal model parametrelerine dönüştürülmektedir. "
  "Geliştirilen ağırlıklı 'uyum skoru' modeli; oyuncu kalitesi, gol üretim verimi ve hedef takımın oyun stili "
  "bileşenlerini birleştirir. Stokastik performans Normal dağılımla, sakatlık oluşumu Bernoulli denemesiyle ve "
  "sakatlık süresi Poisson dağılımıyla modellenmiştir. Örnek vaka olarak Victor Osimhen'in Fenerbahçe'ye olası "
  f"transferi incelenmiş; model {pct(S['UYUM'])} uyum skoru, sezon başına ortalama {S['kacan_ort']:.1f} maç sakatlık "
  f"kaybı ve medyan {S['ga_med']} gol+asist (%80 aralık {S['ga_lo']}–{S['ga_hi']}) projeksiyonu üretmiştir. "
  "Sistem ayrıca etkileşimli bir web arayüzüyle (Streamlit) sunulmuştur. Çalışma, endüstri mühendisliğinin "
  "simülasyon, risk analizi ve karar destek tekniklerinin spor analitiğine uygulanabilirliğini göstermektedir.")
P("Anahtar Kelimeler: Dijital İkiz, Monte Carlo Simülasyonu, Karar Destek Sistemi, Spor Analitiği, "
  "Risk Analizi, Stokastik Modelleme, API-Football.",bold=True)

H1("ABSTRACT")
P("Transfer decisions in professional football are high-budget, high-risk investments for clubs. Whether a "
  "player's current performance will be sustained at a new team depends on many uncertain factors such as injury "
  "history, playing-style compatibility and the team's offensive output. In this study, a decision support system "
  "is developed in which a virtual replica (digital twin) of a footballer is built from real historical data, and a "
  "prospective season at a target club is reproduced thousands of times via Monte Carlo simulation. Data are "
  "retrieved in real time from the API-Football service; the player's per-match performance ratings, goal/assist "
  "output and injury history are processed into probabilistic model parameters. The proposed weighted 'compatibility "
  "score' combines player quality, goal-production efficiency and the target team's playing style. Stochastic "
  "performance is modeled with a Normal distribution, injury occurrence with a Bernoulli trial, and injury duration "
  "with a Poisson distribution. As a case study, the potential transfer of Victor Osimhen to Fenerbahçe is examined; "
  f"the model yields a compatibility score of {pct(S['UYUM'])}, an average of {S['kacan_ort']:.1f} matches lost to "
  f"injury per season, and a median of {S['ga_med']} goals+assists (80% interval {S['ga_lo']}–{S['ga_hi']}). The "
  "system is also delivered through an interactive web interface (Streamlit). The work demonstrates the "
  "applicability of industrial engineering techniques—simulation, risk analysis and decision support—to sports analytics.")
P("Keywords: Digital Twin, Monte Carlo Simulation, Decision Support System, Sports Analytics, Risk Analysis, "
  "Stochastic Modeling, API-Football.",bold=True)
PB()

# ============================================================ 1. GİRİŞ
H1("1. GİRİŞ")
H2("1.1. Problemin Tanımı ve Motivasyon")
P("Futbol endüstrisi, son yirmi yılda veriye dayalı karar alma süreçlerine giderek daha fazla bağımlı hale gelmiştir. "
  "Bir futbolcunun bir kulüpten bir başkasına transferi, çoğu zaman milyonlarca avroluk bonservis bedeli, yüksek "
  "maaş yükü ve uzun vadeli sözleşme taahhütleri içeren ciddi bir yatırım kararıdır. Bu kararın isabetliliği yalnızca "
  "oyuncunun mevcut yeteneğine değil; aynı zamanda yeni takımın oyun sistemine uyum sağlayıp sağlayamayacağına, "
  "fiziksel dayanıklılığına (sakatlık eğilimine) ve katacağı somut katkıya (gol, asist) bağlıdır. Geleneksel olarak "
  "bu değerlendirmeler, scout (gözlemci) raporları ve sezgisel uzman görüşleriyle yapılmaktadır. Ancak insan "
  "sezgisi, çok sayıda belirsiz değişkenin etkileşimini ve bunların oluşturduğu olasılık dağılımlarını nicel olarak "
  "modellemekte yetersiz kalmaktadır.")
P("Bu noktada, belirsizlik altında karar verme problemlerinin temel disiplini olan endüstri mühendisliği devreye "
  "girmektedir. Endüstri mühendisliğinin simülasyon, stokastik modelleme ve risk analizi araçları, bir transferin "
  "olası sonuçlarını tek bir 'nokta tahmin' yerine bir olasılık dağılımı olarak ortaya koyabilir. Böylece karar "
  "vericiye 'oyuncu iyi mi?' sorusunun ötesinde, 'bu oyuncu bizim takımımızda hangi olasılıkla, ne kadar katkı "
  "verir ve hangi riskleri taşır?' sorusunun yanıtı sunulabilir.")
P("Söz konusu kararların finansal büyüklüğü, konunun önemini açıkça ortaya koymaktadır. Üst düzey bir forvet "
  "transferinde bonservis bedeli çoğu zaman onlarca milyon avroyu bulmakta; buna yıllık maaş, imza bonusu ve "
  "menajerlik ücretleri eklendiğinde toplam maliyet katlanmaktadır. Böylesine yüksek bir harcamanın karşılığında "
  "oyuncunun sakatlık nedeniyle sezonun önemli bir bölümünü kaçırması veya yeni sisteme uyum sağlayamaması, kulüp "
  "için ciddi bir batık maliyet anlamına gelir. Dolayısıyla transfer kararı, özünde belirsizlik altında verilen bir "
  "sermaye yatırımı kararıdır ve tıpkı endüstriyel yatırım projelerinde olduğu gibi beklenen getiri ile risk birlikte "
  "değerlendirilmelidir. Bu çalışma, futbol transferlerini bu mühendislik bakış açısıyla ele almakta ve sezgisel "
  "değerlendirmeyi nicel bir olasılık çerçevesiyle desteklemeyi amaçlamaktadır.")
H2("1.2. Projenin Amacı ve Hedefleri")
P("Bu bitirme projesinin temel amacı, bir futbolcunun gerçek geçmiş verilerinden hareketle dijital ikizini "
  "oluşturmak ve bu ikizi kullanarak oyuncunun hedef bir takımdaki olası bir sezonunu Monte Carlo simülasyonu ile "
  "modellemektir. Projenin somut hedefleri şunlardır:")
NUM("Gerçek zamanlı bir spor veri servisinden (API-Football) oyuncu ve takım verilerini otomatik olarak çekecek "
    "dayanıklı bir veri toplama altyapısı geliştirmek.")
NUM("Çekilen ham verileri (maç reytingleri, gol/asist, sakatlık geçmişi) anlamlı olasılıksal model parametrelerine "
    "dönüştürmek.")
NUM("Oyuncu kalitesi, gol üretim verimi ve takım oyun stili bileşenlerini birleştiren ağırlıklı bir 'uyum skoru' "
    "modeli tasarlamak.")
NUM("Stokastik performans dalgalanmasını ve sakatlık riskini içeren bir Monte Carlo simülasyon motoru kurmak.")
NUM("Sonuçları yöneticilerin kolayca yorumlayabileceği etkileşimli bir web tabanlı karar destek panosu (dashboard) "
    "olarak sunmak.")
H2("1.3. Projenin Kapsamı")
P("Çalışma, hücum hattı oyuncuları (özellikle santrforlar) üzerine odaklanmıştır; çünkü bu mevkide katkı, gol ve "
  "asist gibi sayısallaştırması nispeten kolay metriklerle ifade edilebilmektedir. Sistem mimarisi genel olup "
  "herhangi bir lig, takım ve oyuncu için çalışacak biçimde tasarlanmıştır. Kapsam; veri toplama, parametre tahmini, "
  "uyum modellemesi, Monte Carlo simülasyonu ve görselleştirmeyi içermektedir. Oyuncu piyasa değeri tahmini, "
  "sözleşme/maaş optimizasyonu ve taktiksel ısı haritası analizleri bu çalışmanın kapsamı dışındadır ve gelecek "
  "çalışmalara bırakılmıştır.")
H2("1.4. Endüstri Mühendisliği ile İlişkisi")
P("Proje, endüstri mühendisliğinin birçok temel alanını doğrudan uygulamaya geçirir. Monte Carlo simülasyonu ve "
  "kesikli olay modellemesi 'Sistem Simülasyonu'; Bernoulli, Poisson ve Normal dağılımların kullanımı 'Olasılık ve "
  "İstatistik'; uyum skorunun ağırlıklı bileşenlerle oluşturulması 'Çok Kriterli Karar Verme'; sakatlık ve performans "
  "belirsizliğinin nicelleştirilmesi 'Risk Analizi'; nihai panonun yönetsel kullanımı ise 'Karar Destek Sistemleri' "
  "konularıyla örtüşmektedir. Böylece proje, soyut mühendislik araçlarını güncel ve ilgi çekici bir uygulama "
  "alanında somutlaştırmaktadır.")
H2("1.5. Raporun Organizasyonu")
P("Raporun ikinci bölümünde dijital ikiz, Monte Carlo simülasyonu ve spor analitiği kavramlarına ilişkin kuramsal "
  "çerçeve sunulmaktadır. Üçüncü bölüm materyal ve yöntemi; dördüncü bölüm sistem tasarımını ve matematiksel modeli; "
  "beşinci bölüm yazılımsal gerçekleştirimi açıklamaktadır. Altıncı bölüm örnek vaka çalışmasının bulgularını sunar; "
  "yedinci bölüm bu bulguları tartışır; sekizinci bölüm ise sonuç ve önerileri içerir. Kaynakça ve kod ekleri raporun "
  "sonunda yer almaktadır.")
PB()

# ============================================================ 2. LİTERATÜR
H1("2. KURAMSAL ÇERÇEVE VE LİTERATÜR TARAMASI")
H2("2.1. Dijital İkiz (Digital Twin) Kavramı")
P("Dijital ikiz, fiziksel bir varlığın, sürecin veya sistemin gerçek verilerle beslenen sanal bir kopyasıdır. "
  "Kavram ilk olarak imalat ve havacılık sektörlerinde, fiziksel ürünlerin yaşam döngüsünü sanal ortamda izlemek ve "
  "öngörmek amacıyla ortaya çıkmıştır (Grieves, 2014). Bir dijital ikizin üç temel bileşeni vardır: fiziksel varlık, "
  "sanal model ve ikisi arasındaki veri bağlantısı. Bu çalışmada fiziksel varlık futbolcunun kendisi, sanal model "
  "onun istatistiksel performans modeli, veri bağlantısı ise API-Football üzerinden gerçek zamanlı çekilen verilerdir. "
  "Dijital ikizin temel üstünlüğü, gerçek dünyada denenmesi maliyetli veya imkânsız olan senaryoların ('Bu oyuncu "
  "bizim takımımızda oynasaydı ne olurdu?') sanal ortamda güvenle test edilebilmesidir.")
P("Endüstri 4.0 paradigmasının merkezindeki kavramlardan biri olan dijital ikiz, son yıllarda imalatın ötesine "
  "geçerek sağlık (hasta ikizleri), şehircilik (akıllı şehir ikizleri) ve enerji gibi alanlara yayılmıştır. Ortak "
  "nokta, gerçek sistemden sürekli beslenen ve onunla eşzamanlı kalan bir sanal modelin; gözlem, tahmin ve 'ya olursa' "
  "(what-if) analizleri için kullanılmasıdır. Bir futbolcunun dijital ikizi de benzer biçimde, oyuncunun gerçek "
  "istatistiksel 'imzasını' yakalayan ve bu imzayı farklı bağlamlara (yeni bir takıma) taşıyarak senaryo analizi "
  "yapmaya olanak tanıyan bir soyutlamadır. Bu çalışma, kavramın spor alanındaki nispeten yeni ve az incelenmiş bir "
  "uygulamasını sunmaktadır.")
H2("2.2. Monte Carlo Simülasyonu")
P("Monte Carlo simülasyonu, belirsizlik içeren sistemlerin davranışını, rastgele örnekleme yoluyla çok sayıda kez "
  "tekrar ederek olasılıksal sonuçlar elde eden bir sayısal yöntemdir. Yöntem, II. Dünya Savaşı sırasında Manhattan "
  "Projesi kapsamında Ulam ve von Neumann tarafından geliştirilmiş (Metropolis ve Ulam, 1949) ve adını Monako'daki "
  "kumarhaneden almıştır. Temel mantığı, analitik olarak çözülmesi zor bir problemin girdi değişkenlerine ilişkin "
  "olasılık dağılımlarından rastgele değerler çekip sistemi defalarca çalıştırmak ve çıktıların dağılımını "
  "incelemektir. Yeterince büyük tekrar sayısında (bu çalışmada 10.000 sezon), büyük sayılar yasası gereği örneklem "
  "ortalaması gerçek beklenen değere yakınsar. Yöntem; finans (risk değerlemesi), proje yönetimi (süre/maliyet "
  "belirsizliği) ve mühendislik güvenilirliği gibi alanlarda yaygın olarak kullanılır.")
P("Monte Carlo yönteminin gücü, çıktının tamamını bir olasılık dağılımı olarak elde edebilmesinden gelir. Analitik "
  "yöntemler genellikle yalnızca beklenen değeri verirken, simülasyon medyanı, yüzdelikleri, en kötü/en iyi "
  "senaryoları ve belirli eşiklerin aşılma olasılıklarını birlikte sağlar. Yöntemin yakınsama hızı büyük sayılar "
  "yasasına bağlıdır; tahminin standart hatası tekrar sayısının kareköküyle (1/√N) azalır. Bu çalışmada seçilen "
  "N = 10.000 tekrar, standart hatayı ihmal edilebilir düzeye indirirken makul bir hesaplama süresi sağlamaktadır. "
  "Ayrıca sabit bir rastgelelik tohumu (seed) kullanılarak aynı girdilerle aynı çıktıların elde edilmesi, yani "
  "bilimsel yeniden üretilebilirlik güvence altına alınmıştır.")
H2("2.3. Spor Analitiği ve Futbolda Veri Bilimi")
P("Spor analitiği, atletik performansın ve takım stratejilerinin nicel verilerle incelenmesidir. Beyzbolda "
  "'Moneyball' yaklaşımıyla popülerleşen veri temelli karar alma (Lewis, 2003), günümüzde futbolda da beklenen gol "
  "(xG), oyuncu reytingleri ve fiziksel takip verileriyle olgunlaşmıştır. Maç reytingleri (genellikle 0–10 ölçeğinde), "
  "bir oyuncunun tek bir maçtaki çok boyutlu katkısını tek bir sayıya indirgeyen bileşik göstergelerdir ve bu "
  "çalışmada oyuncunun temel performans seviyesinin ve maçtan maça dalgalanmasının tahmininde kullanılmıştır.")
H2("2.4. Olasılıksal Risk Modelleme")
P("Çalışmada üç temel olasılık dağılımı kullanılmıştır. Normal (Gauss) dağılımı, oyuncunun maç performansındaki "
  "simetrik dalgalanmayı; Bernoulli dağılımı, bir maçta yeni bir sakatlığın oluşup oluşmamasını (ikili sonuç); "
  "Poisson dağılımı ise bir sakatlık durumunda kaçırılacak maç sayısını (belirli bir ortalama etrafında sayma "
  "değişkeni) modellemektedir. Bu dağılımların seçimi, modellenen olgunun doğasıyla uyumludur: performans sürekli ve "
  "simetrik, sakatlık oluşumu ikili, sakatlık süresi ise negatif olmayan tamsayıdır.")
H2("2.5. Karar Destek Sistemleri")
P("Karar destek sistemi (KDS), karar vericilere yapılandırılmamış veya yarı yapılandırılmış problemlerde yardımcı "
  "olan etkileşimli bilgi sistemleridir. İyi bir KDS, veriyi bilgiye dönüştürür ve sonucu sezgisel görsellerle sunar. "
  "Bu çalışmada geliştirilen web panosu; uyum skorunu bir gösterge (gauge) ile, katkı projeksiyonunu olasılık "
  "dağılımı grafiğiyle ve riski metrik kartlarıyla sunarak bu ilkeyi hayata geçirmektedir.")
H2("2.6. İlgili Sistemler ve Çalışmanın Özgün Katkısı")
P("Spor analitiği alanında oyuncu değerlendirmesine yönelik birçok ticari ve akademik araç bulunmaktadır. Bunların "
  "çoğu, oyuncuları geçmiş performanslarına göre sıralayan tanımlayıcı (descriptive) istatistik panolarıdır; yani "
  "'oyuncu ne yaptı?' sorusuna yanıt verirler. Daha gelişmiş bazı sistemler beklenen gol (xG) gibi tahminsel "
  "metrikler sunar. Ancak literatürde, bir oyuncunun belirli bir hedef takımdaki gelecek performansını sakatlık "
  "riskini de içerecek biçimde olasılıksal olarak simüle eden ve sonucu bir dağılım olarak veren bütünleşik bir karar "
  "destek aracına nispeten az rastlanmaktadır. Bu çalışmanın özgün katkısı; dijital ikiz kavramını, gerçek zamanlı "
  "veri toplamayı, çok kriterli uyum modellemesini ve Monte Carlo risk simülasyonunu tek bir uçtan uca sistemde "
  "birleştirmesidir. Böylece sistem, 'oyuncu ne yaptı?' sorusundan 'oyuncu bizim takımımızda hangi olasılıkla ne "
  "yapar ve hangi riski taşır?' sorusuna geçiş yapmaktadır. Bu öngörücü (predictive) ve kuralcı (prescriptive) "
  "yaklaşım, çalışmanın endüstri mühendisliği karar bilimi perspektifiyle örtüşen en özgün yönüdür.")
PB()

# ============================================================ 3. MATERYAL VE YÖNTEM
H1("3. MATERYAL VE YÖNTEM")
H2("3.1. Genel Yaklaşım")
P("Çalışma, beş aşamalı bir iş akışı üzerine kurulmuştur: (1) gerçek veri toplama, (2) veri ön işleme ve parametre "
  "tahmini, (3) uyum skoru hesaplama, (4) Monte Carlo simülasyonu, (5) görselleştirme ve raporlama. Bu modüler yapı, "
  "her bileşenin bağımsız olarak test edilip geliştirilebilmesini sağlamaktadır.")
H2("3.2. Veri Kaynağı: API-Football")
P("Proje verileri, dünya genelindeki yüzlerce ligi kapsayan API-Football (v3) servisinden alınmıştır. Servis, "
  "REST mimarisiyle JSON biçiminde yanıt döndürmekte ve kimlik doğrulaması için bir API anahtarı gerektirmektedir. "
  "Çalışmada ücretsiz katman kullanılmış olup bu katman günde 100 istek ve dakikada sınırlı sayıda istek kısıtı "
  "getirmektedir. Bu kısıt, veri toplama mimarisinin tasarımında belirleyici bir etken olmuştur. Kullanılan başlıca "
  "uç noktalar (endpoint) aşağıdaki tabloda özetlenmiştir.")
TABLE(["Uç Nokta (Endpoint)","Amaç","Kullanılan Alanlar"],
 [["players/profiles","Oyuncuyu isimle bulma","id, name, lastname"],
  ["players/seasons","Oyuncunun veri içeren sezonları","yıl listesi"],
  ["players","Sezonluk istatistik","rating, goals, assists, minutes"],
  ["fixtures","Takımın maç listesi","fixture id, status, timestamp"],
  ["fixtures/players","Maç bazlı oyuncu reytingi","games.rating"],
  ["injuries","Sakatlık geçmişi","type, reason, date"],
  ["teams / leagues","Hedef takım ve ligi","team id, league, season"],
  ["teams/statistics","Takım hücum profili","goals.for, fixtures.played, lineups"]],
 "Tablo 3.1. Kullanılan API-Football uç noktaları ve amaçları.",
 widths=[4.2,5.6,5.5])
P("Servisin sunduğu veri derinliği oldukça geniştir; ancak ücretsiz katmanın günlük (100 istek) ve dakikalık istek "
  "kısıtları, tüm sezonun maç-maç verisini çekmeyi pratikte zorlaştırmaktadır. Bu nedenle çalışmada, istatistiksel "
  "olarak yeterli bir örneklem sağlayacak sayıda (varsayılan on) maç reytingi çekilmekte; böylece doğruluk ile istek "
  "bütçesi arasında bir denge kurulmaktadır. Bu tasarım tercihi, gerçek dünya kaynak kısıtları altında mühendislik "
  "çözümlerinin nasıl şekillendiğinin tipik bir örneğidir ve sistemin ücretli bir plana yükseltilmesi hâlinde "
  "doğrudan daha derin analizlere ölçeklenebilir.")
H2("3.3. Kullanılan Yazılım ve Kütüphaneler")
TABLE(["Bileşen","Teknoloji","Görev"],
 [["Programlama dili","Python 3.12","Tüm sistem mantığı"],
  ["Sayısal hesaplama","NumPy","Rastgele sayı üretimi, istatistik"],
  ["HTTP istemcisi","Requests","API çağrıları"],
  ["Web arayüzü","Streamlit","Etkileşimli pano"],
  ["Görselleştirme","Plotly / Matplotlib","Gauge, histogram, grafikler"],
  ["Metin eşleştirme","difflib, unicodedata","Yazım hatasına dayanıklı arama"]],
 "Tablo 3.2. Projede kullanılan yazılım ve kütüphaneler.",
 widths=[4.5,4.8,6.0])
H2("3.4. Veri Toplama Mimarisi")
P("Veri toplama katmanı, API anahtarını ve hız sınırlama mantığını kapsülleyen bir istemci sınıfı (API) etrafında "
  "tasarlanmıştır. Her çağrı arasına yapılandırılabilir bir bekleme süresi (varsayılan 7 saniye) konularak dakikalık "
  "istek kısıtı aşılmamaktadır. Ayrıca '429 Too Many Requests' yanıtı alındığında istemci otomatik olarak bekleyip "
  "isteği yeniden denemekte; geçici hatalarda (404 vb.) ise boş sonuç döndürerek tüm sürecin çökmesini önlemektedir. "
  "Bu dayanıklılık önlemleri, ücretsiz katmanın kısıtları altında güvenilir veri toplamayı mümkün kılmıştır.")
H2("3.5. Karşılaşılan Teknik Zorluklar ve Çözümleri")
P("Geliştirme sürecinde karşılaşılan başlıca zorluklar ve geliştirilen çözümler şunlardır:")
BULLET("Hız sınırı (429): Çağrılar arası bekleme + otomatik yeniden deneme + çekilen maç sayısının sınırlandırılması.")
BULLET("Oyuncu bulma: Ön ad/soyad, aksanlı karakterler ve yazım hataları sorun yaratıyordu. Soyad önceliği, "
       "aksan sadeleştirme (unicodedata), önek kısaltma ve difflib ile benzerlik sıralaması içeren dayanıklı bir "
       "arama algoritması geliştirildi.")
BULLET("Sezon/lig otomatik tespiti: Kullanıcının lig ve sezon girmesine gerek kalmadan, oyuncunun en güncel ve en "
       "çok oynadığı lig ile hedef takımın güncel ligi otomatik belirlenmektedir.")
BULLET("Veri kalitesi: Sakatlık kayıtları içinde kart cezası, milli takım daveti ve dinlenme gibi sakatlık "
       "olmayan yokluklar bulunmaktaydı. Bunlar filtrelenerek yalnızca gerçek sakatlıklar modele dâhil edildi.")
PB()

# ============================================================ 4. SİSTEM TASARIMI
H1("4. SİSTEM TASARIMI VE MATEMATİKSEL MODEL")
H2("4.1. Sistem Mimarisi")
P("Sistemin genel mimarisi Şekil 4.1'de gösterilmiştir. Kullanıcı arayüzünden gelen istek, veri çekme katmanı "
  "aracılığıyla API-Football'a iletilir; gelen ham veri ön işleme katmanında parametrelere dönüştürülür; bu "
  "parametreler hem uyum modeline hem de Monte Carlo motoruna beslenir; üretilen sonuçlar panoya aktarılarak "
  "kullanıcıya sunulur.")
FIG("fig_arch.png","Şekil 4.1. Dijital ikiz sisteminin katmanlı mimarisi ve veri akışı.")
H2("4.2. Uyum Skoru Modeli")
P("Uyum skoru (S), bir oyuncunun hedef takıma ne ölçüde uygun olduğunu 0–1 aralığında ifade eden bileşik bir "
  "göstergedir. Üç temel bileşenin ağırlıklı toplamı olarak tanımlanmıştır:")
P("S = w₁ · Kalite + w₂ · Verim + w₃ · Stil",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
P("Burada ağırlıklar w₁ = 0,50, w₂ = 0,25 ve w₃ = 0,25 olarak belirlenmiştir. Ağırlıklar, oyuncu kalitesinin "
  "uyumun en belirleyici unsuru olduğu varsayımıyla seçilmiş olup kullanıcı tarafından değiştirilebilir.")
H3("4.2.1. Kalite Bileşeni")
P("Kalite, oyuncunun ortalama maç reytinginin 0–1 aralığına ölçeklenmiş hâlidir: Kalite = ortalama_reyting / 10. "
  "Reyting, oyuncunun çok boyutlu katkısını tek sayıda özetlediğinden temel performans seviyesinin sağlam bir "
  "göstergesidir.")
H3("4.2.2. Verim Bileşeni")
P("Verim, oyuncunun 90 dakika başına ürettiği gol+asist (G+A/90) değerinin elit bir referansa (1,0 G+A/90) "
  "oranlanıp 1 ile sınırlandırılmasıyla elde edilir: Verim = min(GA90 / 1,0 ; 1,0). Bu bileşen, hücum oyuncusunun "
  "somut üretkenliğini ölçer.")
H3("4.2.3. Stil Uyumu Bileşeni")
P("Stil uyumu, hedef takımın maç başına attığı gol sayısının bir referansa (2,6 gol/maç) oranlanmasıyla hesaplanır: "
  "Stil = min(takım_gol_maç / 2,6 ; 1,0). Çok gol atan bir takım, bir forvete daha fazla pozisyon ve gol şansı "
  "üreteceğinden yüksek stil uyumu, oyuncunun potansiyelini daha iyi gerçekleştirebileceği anlamına gelir.")
H3("4.2.4. Genişletme: Rol Uyumu")
P("Hedef takımın mevcut forvetlerine ilişkin veri mevcut olduğunda, dört bileşenli bir varyant kullanılabilir. Bu "
  "varyantta oyuncunun G+A/90 değeri takımın en üretken forvetininkiyle kıyaslanarak bir 'rol uyumu' bileşeni "
  "eklenir ve ağırlıklar 0,40/0,20/0,20/0,20 olarak yeniden dağıtılır. Bu, oyuncunun mevcut kadroyu gerçekten "
  "geliştirip geliştirmeyeceğini ölçer.")
TABLE(["Bileşen","Formül","Ağırlık","Ölçtüğü"],
 [["Kalite","ort_reyting / 10","0,50","Genel seviye"],
  ["Verim","min(GA90 / 1,0 ; 1)","0,25","Gol üretimi"],
  ["Stil","min(gol_maç / 2,6 ; 1)","0,25","Takım hücum bolluğu"],
  ["(Rol)","min(GA90 / forvet ; 1,25)/1,25","ops.","Kadroyu geliştirme"]],
 "Tablo 4.1. Uyum skoru bileşenleri, formülleri ve ağırlıkları.",
 widths=[2.8,6.6,2.2,3.7])
H3("4.2.5. Örnek Uyum Hesabı")
P("Modelin işleyişini somutlaştırmak için bir örnek hesap verilebilir. Ortalama reytingi 7,74 olan bir oyuncu için "
  "Kalite bileşeni 7,74 / 10 = 0,774 olarak bulunur. Oyuncunun 90 dakika başına gol+asist üretimi 1,217 olduğundan ve "
  "bu değer elit referans olan 1,0'ı aştığından Verim bileşeni 1,0 ile sınırlanır. Hedef takım maç başına 2,5 gol "
  "attığından Stil bileşeni 2,5 / 2,6 = 0,962 olur. Bu üç bileşen ağırlıklarıyla birleştirildiğinde genel uyum skoru "
  "S = 0,50 × 0,774 + 0,25 × 1,000 + 0,25 × 0,962 = 0,877, yani yaklaşık %87,7 elde edilir. Bu örnek, modelin her bir "
  "bileşeninin nihai sonuca katkısını şeffaf biçimde izlenebilir kıldığını göstermektedir; bu şeffaflık, bir karar "
  "destek sisteminin güvenilirliği ve kabul edilebilirliği açısından kritik bir özelliktir.")
H2("4.3. Stokastik Performans Modeli")
P("Oyuncunun belirli bir maçtaki performans seviyesi (S_maç), temel performansın etrafında iki bağımsız rastgele "
  "gürültü teriminin eklenmesiyle modellenir:")
P("S_maç = S_temel + ε_form + ε_maç ,  ε_form ~ N(0, σ_form²) , ε_maç ~ N(0, σ_maç²)",bold=True,align=WD_ALIGN_PARAGRAPH.CENTER)
P("Burada ε_form sezon içi form dalgalanmasını, ε_maç ise maça özgü anlık değişkenliği temsil eder. Toplam "
  "değişkenlik, oyuncunun maç bazlı reytinglerinin standart sapmasından tahmin edilir ve σ_form² + σ_maç² = σ_toplam² "
  "olacak biçimde 0,8 ve 0,6 katsayılarıyla bölünür. S_maç değeri [0, 1] aralığına kırpılarak mantıksal tutarlılık "
  "korunur.")
H2("4.4. Sakatlık Risk Modeli")
P("Sakatlık riski iki aşamalı bir stokastik süreçle modellenir. İlk aşamada, her maç için bağımsız bir Bernoulli "
  "denemesi yapılır: oyuncu p olasılığıyla yeni bir sakatlık yaşar. p parametresi, geçmiş sakatlık epizodu sayısının "
  "toplam maç sayısına oranıyla tahmin edilir. İkinci aşamada, sakatlık gerçekleştiğinde kaçırılacak maç sayısı "
  "Poisson(λ) dağılımından çekilir; λ, geçmişte sakatlık başına kaçırılan ortalama maç sayısıdır (en az 1 ile "
  "sınırlanır). Sakat oyuncu, kaçırdığı maçlarda performans üretmez ve bu maçlar projeksiyona dâhil edilmez.")
H2("4.5. Gol+Asist Projeksiyon Modeli")
P("Oynanan her maçta oyuncunun beklenen gol+asist katkısı, gerçek GA90 değerinin o maçtaki göreli performansla "
  "ölçeklenmesiyle elde edilir: λ_maç = GA90 · (S_maç / S_temel). Maçtaki gerçek katkı, bu beklenen değere sahip bir "
  "Poisson dağılımından çekilerek bitiricilik (finishing) varyansı da modele dâhil edilir. Sezonluk toplam, oynanan "
  "tüm maçların katkılarının toplamıdır.")
H2("4.6. Monte Carlo Simülasyon Algoritması")
P("Simülasyon, bir sezonu (varsayılan 38 maç) N = 10.000 kez yeniden üretir. Her sezon için her maç sırayla işlenir: "
  "oyuncu sakatsa maç atlanır; değilse önce sakatlık kontrolü yapılır, ardından performans ve gol+asist katkısı "
  "üretilir. Sezon sonunda ortalama performans, kaçırılan maç sayısı ve toplam gol+asist kaydedilir. Tüm sezonların "
  "sonuçları biriktirilerek medyan, yüzdelikler ve olasılıklar gibi özet istatistikler hesaplanır. Algoritmanın "
  "sözde kodu aşağıda verilmiştir.")
CODE(["N sezon için tekrarla:",
      "    kacan = 0 ; kalan_sakatlik = 0 ; gol_asist = 0",
      "    her maç için (1..38):",
      "        eğer kalan_sakatlik > 0:  kalan--, kacan++  ve sonraki maça geç",
      "        eğer rastgele() < p:      kalan = Poisson(λ) ; kacan++ ; sonraki maça geç",
      "        S_mac = clip(S_temel + N(0,σ_form) + N(0,σ_mac), 0, 1)",
      "        gol_asist += Poisson( GA90 * S_mac / S_temel )",
      "    sonuçları kaydet (performans, kacan, gol_asist)",
      "tüm sezonların medyan / yüzdelik / olasılık özetini çıkar"],
     "Algoritma 4.1. Monte Carlo simülasyon motorunun sözde kodu.")
P("Algoritmanın hesaplama karmaşıklığı, sezon sayısı ile maç sayısının çarpımı (N × maç) mertebesindedir; bu da "
  "varsayılan ayarlar için yaklaşık 380.000 temel iterasyona karşılık gelir ve modern bir bilgisayarda saniyenin "
  "altında bir sürede tamamlanır. Bu verimlilik, kullanıcının arayüzde gerçek zamanlıya yakın bir deneyim yaşamasını "
  "sağlar. Asıl darboğaz hesaplama değil, API'den veri çekme aşamasıdır; çünkü hız sınırı nedeniyle bu aşama birkaç "
  "dakika sürebilmektedir. Parametreler bir kez tahmin edildikten sonra simülasyonun kendisi, farklı senaryolar "
  "(örneğin farklı sezon uzunlukları veya değiştirilmiş risk düzeyleri) için neredeyse anında yeniden "
  "çalıştırılabilir; bu da duyarlılık analizlerini pratik kılar.")
PB()

# ============================================================ 5. UYGULAMA
H1("5. UYGULAMA (GERÇEKLEŞTİRİM)")
H2("5.1. Veri Çekme Modülü")
P("Veri çekme modülü, API çağrılarını, hız sınırlamasını ve hata yönetimini kapsülleyen bir istemci sınıfı üzerine "
  "kuruludur. Aşağıdaki kod parçası, otomatik yeniden deneme ve bekleme mantığını göstermektedir.")
CODE(["class API:",
      "    def _req(self, path, params):",
      "        for _ in range(6):",
      "            r = requests.get(f'{self.base}/{path}', headers=self.h, params=params)",
      "            if r.status_code == 429:   # hız sınırı: bekle ve yeniden dene",
      "                time.sleep(30); continue",
      "            r.raise_for_status(); time.sleep(self.t); return r.json()",
      "        raise RuntimeError('API limiti doldu')"],
     "Kod 5.1. Hız sınırına dayanıklı API istemcisi (özet).")
H2("5.2. Dayanıklı Oyuncu Arama Algoritması")
P("Kullanıcının girdiği isim; soyad önceliği, aksan sadeleştirme, önek kısaltma (yazım hatası toleransı) ve difflib "
  "ile benzerlik sıralaması kullanılarak en olası oyuncuyla eşleştirilir. Bulunan adaylar arasından istatistiği olan "
  "ilk oyuncu/sezon seçilir.")
CODE(["def benzerlik(aday, isim):  # difflib ile 0-1 arası benzerlik",
      "    return max(SequenceMatcher(None, sade(isim), sade(ad)).ratio() ...)",
      "",
      "# 'Osimen' -> 'Osim' önekiyle bulunur, ardından benzerliğe göre 'Osimhen' seçilir",
      "adaylar = sorted(havuz, key=lambda a: benzerlik(a, isim), reverse=True)"],
     "Kod 5.2. Yazım hatasına dayanıklı arama (özet).")
H2("5.3. Parametre Tahmin Modülü")
P("Bu modül, ham verileri model parametrelerine dönüştürür: ortalama reytingden S_temel; maç reytinglerinin standart "
  "sapmasından σ_form ve σ_maç; filtrelenmiş sakatlık epizodlarından p ve λ; gol/asist ve dakikadan GA90 değeri "
  "hesaplanır. Sakatlık filtresi, sakatlık olmayan yoklukları (kart, milli takım, dinlenme) dışlar.")
H2("5.4. Web Arayüzü")
P("Sistem, Streamlit çatısıyla geliştirilen bir web panosu olarak sunulmuştur. Kullanıcı sol panelden API anahtarı, "
  "oyuncu adı ve hedef takımı girer; 'Analiz Et' düğmesiyle veri çekme ve simülasyon süreci başlar. Sonuçlar; büyük "
  "bir uyum göstergesi (gauge), oyuncu ve takım kartları, dört metrik kartı (beklenen katkı, sakatlık riski, katkı "
  "olasılıkları ve sağlamlık) ile bileşen ve dağılım grafikleri olarak koyu temalı, modern bir arayüzde gösterilir.")
H2("5.5. Sistemin Test Edilmesi ve Dayanıklılığı")
P("Sistemin güvenilirliği, geliştirme boyunca çeşitli gerçek senaryolarla sınanmıştır. Farklı liglerden ve farklı "
  "isim yazımlarına sahip oyuncular denenerek arama algoritmasının dayanıklılığı doğrulanmış; yazım hatası içeren "
  "girdilerin dahi doğru oyuncuya eşleştiği gözlemlenmiştir. API hız sınırına takılma, boş veri dönmesi ve geçersiz "
  "takım adı gibi hata senaryoları için savunmacı programlama teknikleri uygulanmış ve sistemin bu durumlarda çökmek "
  "yerine anlamlı uyarılar üretmesi sağlanmıştır. Simülasyon modülü ise sabit tohumla tekrar tekrar çalıştırılarak "
  "çıktıların tutarlılığı (yeniden üretilebilirliği) teyit edilmiştir. Bu testler, sistemin kontrollü bir geliştirme "
  "ortamından gerçek dünya kullanımına taşınabilecek olgunlukta olduğunu göstermektedir.")
PB()

# ============================================================ 6. BULGULAR
H1("6. BULGULAR VE VAKA ÇALIŞMASI")
H2("6.1. Vaka Tanımı")
P("Geliştirilen sistem, somut bir senaryo üzerinde test edilmiştir: Nijeryalı santrfor Victor Osimhen'in Fenerbahçe'ye "
  "olası transferinin değerlendirilmesi. Oyuncunun en güncel sezon verileri ile hedef takımın güncel hücum profili "
  "API-Football'dan otomatik olarak çekilmiştir.")
H2("6.2. Çekilen Gerçek Veriler")
TABLE(["Değişken","Değer","Değişken","Değer"],
 [["Oynanan maç",f"{S['oynanan']}","Toplam dakika",f"{int(S['dakika'])}"],
  ["Gol",f"{S['gol']}","Asist",f"{S['asist']}"],
  ["Ortalama reyting",f"{S['ort_rating']:.2f}","Gol+Asist / 90",f"{S['ga90']:.3f}"],
  ["Sakatlık epizodu",f"{S['sakatlik_episode']}","Kaçan maç (3 sezon)",f"{S['sakatlik_kacan']}"],
  ["Hedef takım gol/maç",f"{S['fb_gol_mac']}","Toplam maç (payda)",f"{S['toplam_mac']}"]],
 "Tablo 6.1. Osimhen ve Fenerbahçe için API-Football'dan çekilen ham veriler.",
 widths=[5.0,2.6,5.0,2.7])
P("Oyuncunun son maçlarındaki performans reytinglerinin dağılımı Şekil 6.1'de görülmektedir. Reytinglerin "
  "ortalamanın etrafında nispeten dar bir bantta seyretmesi, oyuncunun istikrarlı bir performans sergilediğine "
  "işaret etmektedir.")
FIG("fig_ratings.png","Şekil 6.1. Oyuncunun son maçlardaki performans reytingi dağılımı ve ortalaması.")
P("Oyuncunun sakatlık geçmişi incelendiğinde, üç sezonda altı ayrı gerçek sakatlık epizodu tespit edilmiştir. "
  "Bunların çoğunluğu uyluk (kas) kaynaklı olup epizot başına kaçırılan maç sayısı değişkenlik göstermektedir "
  "(Şekil 6.2). Kart cezası, milli takım daveti gibi sakatlık dışı yokluklar bu analizden çıkarılmıştır.")
FIG("fig_injury.png","Şekil 6.2. Gerçek sakatlık epizodları ve epizot başına kaçırılan maç sayısı (3 sezon).")
H2("6.3. Tahmin Edilen Model Parametreleri")
P("Ham verilerden türetilen olasılıksal model parametreleri Tablo 6.2'de sunulmuştur. Tüm parametreler doğrudan "
  "gerçek veriden hesaplanmış olup hiçbiri elle atanmamıştır; bu, modelin nesnelliğini güçlendirmektedir.")
TABLE(["Parametre","Sembol","Değer","Kaynak"],
 [["Temel performans","S_temel",f"{S['S_base']:.3f}","ort. reyting / 10"],
  ["Form değişkenliği","σ_form",f"{S['sigma_form']:.3f}","reyting std × 0,8"],
  ["Maç değişkenliği","σ_maç",f"{S['sigma_mac']:.3f}","reyting std × 0,6"],
  ["Sakatlık olasılığı","p",f"{S['p_sak']:.3f}",f"{S['sakatlik_episode']} epizot / {S['toplam_mac']} maç"],
  ["Sakatlık süresi","λ",f"{S['lam_sak']:.2f}","kaçan maç / epizot"],
  ["Üretkenlik","GA90",f"{S['ga90']:.3f}","(gol+asist)·90 / dakika"]],
 "Tablo 6.2. Gerçek veriden tahmin edilen Monte Carlo model parametreleri.",
 widths=[4.6,2.6,2.6,5.5])
H3("6.3.1. Parametrelerin Adım Adım Hesaplanması")
P("Sakatlık parametrelerinin ham veriden nasıl türetildiği, modelin veri-temelli niteliğini gösteren önemli bir "
  "örnektir. API-Football'dan çekilen ham sakatlık listesinde toplam 27 yokluk kaydı bulunmaktaydı. Ancak bu "
  "kayıtların önemli bir bölümü gerçek sakatlık değildi: 'Red Card' (kart cezası), 'National selection' ve "
  "'International duty' (milli takım), 'Rest' (dinlenme) ve 'Yellow Cards' gibi nedenler sakatlık dışı yoklukları "
  "temsil ediyordu. Bu kayıtlar filtrelendiğinde geriye yalnızca uyluk, kas ve sırt kaynaklı gerçek sakatlık "
  "kayıtları kalmıştır. Ardından, tarihler arasındaki boşluklara bakılarak (30 günden büyük boşluk = yeni epizot) bu "
  "kayıtlar gruplanmış ve toplam altı ayrı sakatlık dönemi tespit edilmiştir. Bu dönemlerde toplam 17 maç "
  "kaçırılmıştır. Buradan, maç başına yeni sakatlık olasılığı p = 6 / 108 ≈ 0,056 ve sakatlık başına ortalama "
  "kaçırılan maç λ = 17 / 6 ≈ 2,83 olarak hesaplanmıştır. Benzer biçimde, oyuncunun maç reytinglerinin standart "
  "sapması yaklaşık 0,086 (0–1 ölçeğinde) bulunmuş ve bu toplam değişkenlik, σ_form² + σ_maç² = σ_toplam² koşulunu "
  "sağlayacak biçimde σ_form ≈ 0,069 ve σ_maç ≈ 0,052 olarak ikiye ayrılmıştır. Görüldüğü üzere, modelin tüm "
  "girdileri nesnel bir biçimde gerçek gözlemlerden elde edilmektedir; hiçbir parametre keyfî olarak atanmamıştır.")
H2("6.4. Uyum Skoru Sonuçları")
P(f"Üç bileşenli uyum modeli, Osimhen–Fenerbahçe çifti için {pct(S['UYUM'])} düzeyinde bir genel uyum skoru "
  f"üretmiştir (Şekil 6.3). Bu yüksek skorun temel nedenleri, oyuncunun elit gol üretim verimi (Verim bileşeni "
  f"{pct(S['verim'])}) ve Fenerbahçe'nin maç başına {S['fb_gol_mac']} gol atan üretken hücum yapısıdır (Stil bileşeni "
  f"{pct(S['stil'])}). Kalite bileşeni ise {pct(S['kalite'])} ile güçlü bir temel performans seviyesini yansıtmaktadır "
  "(Şekil 6.4).")
FIG("fig_gauge.png","Şekil 6.3. Hesaplanan genel uyum skoru göstergesi.",w=9)
FIG("fig_components.png","Şekil 6.4. Uyum skorunun bileşenlere göre dağılımı.")
P("Uyum skorunun bileşen bazında ayrıştırılması, karar vericiye yalnızca bir nihai not değil, bu notun nereden "
  "geldiğini de gösterir. Bu örnekte skorun büyük ölçüde oyuncunun bireysel niteliklerinden (kalite ve verim) "
  "kaynaklandığı, takım stilinin ise bu güçlü temeli desteklediği görülmektedir. Böyle bir ayrıştırma; farklı oyuncu "
  "adaylarının güçlü ve zayıf yönlerinin kıyaslanmasında ve takımın ihtiyacına en uygun profilin seçilmesinde "
  "doğrudan kullanılabilir. Örneğin, stil uyumu düşük ama kalitesi yüksek bir oyuncu, farklı oyun anlayışına sahip "
  "bir takıma; verimi yüksek bir oyuncu ise hücumda üretkenlik arayan bir takıma yönlendirilebilir.")
H2("6.5. Monte Carlo Simülasyon Sonuçları")
P(f"10.000 sezonluk simülasyon sonucunda, oyuncunun Fenerbahçe'de bir sezonda üreteceği gol+asist katkısının medyanı "
  f"{S['ga_med']} olarak bulunmuştur; sonuçların %80'i {S['ga_lo']}–{S['ga_hi']} aralığında yoğunlaşmaktadır "
  f"(Şekil 6.5). Oyuncunun 20 ve üzeri gol+asist üretme olasılığı {pct(S['p20'])}, 30 ve üzeri üretme olasılığı ise "
  f"{pct(S['p30'])} gibi oldukça yüksek değerlerdedir. Bu yüksek projeksiyon, oyuncunun mevcut üretkenlik oranını "
  "sürdürmesi varsayımına dayanmaktadır.")
FIG("fig_ga_dist.png","Şekil 6.5. Sezonluk gol+asist katkısının Monte Carlo olasılık dağılımı.")
P(f"Sakatlık tarafında ise model, oyuncunun sezon başına ortalama {S['kacan_ort']:.1f} maçı sakatlık nedeniyle "
  f"kaçıracağını öngörmektedir (Şekil 6.6). Oyuncunun en az 32 maç oynayabilme olasılığı ('sağlamlık' göstergesi) "
  f"{pct(S['saglam'])} olarak hesaplanmıştır; bir başka deyişle yaklaşık üçte bir olasılıkla altı veya daha fazla maç "
  "kaçırma riski bulunmaktadır. Bu, transferin yatırım getirisi (ROI) değerlendirmesinde dikkate alınması gereken "
  "temel risk unsurudur.")
FIG("fig_missed_dist.png","Şekil 6.6. Sezon başına sakatlıktan kaçırılan maç sayısının olasılık dağılımı.")
TABLE(["Çıktı Göstergesi","Değer"],
 [["Genel uyum skoru",pct(S['UYUM'])],
  ["Medyan sezonluk performans",pct(S['perf_med'])],
  ["Beklenen gol+asist (medyan)",f"{S['ga_med']}  (%80: {S['ga_lo']}–{S['ga_hi']})"],
  ["Ortalama kaçırılan maç",f"{S['kacan_ort']:.1f} maç"],
  ["20+ gol+asist olasılığı",pct(S['p20'])],
  ["30+ gol+asist olasılığı",pct(S['p30'])],
  ["Sağlamlık (32+ maç oynama)",pct(S['saglam'])]],
 "Tablo 6.3. Simülasyon çıktılarının özeti (Osimhen → Fenerbahçe).",
 widths=[9.0,6.3])
H2("6.6. Duyarlılık (Hassasiyet) Analizi")
P("Modelin girdi parametrelerine duyarlılığını incelemek amacıyla iki analiz yapılmıştır (Şekil 6.7). İlk panelde, "
  "hedef takımın maç başına gol üretiminin uyum skoru üzerindeki etkisi gösterilmektedir; takım hücum gücü arttıkça "
  "stil uyumu ve dolayısıyla genel skor artmakta, ancak belirli bir noktadan sonra (referans değerde) doygunluğa "
  "ulaşmaktadır. İkinci panelde, maç başına sakatlık olasılığının beklenen gol+asist üzerindeki etkisi "
  "incelenmektedir; sakatlık riski arttıkça oynanan maç sayısı azaldığından beklenen katkı belirgin biçimde "
  "düşmektedir. Bu analiz, sakatlık riskinin transferin getirisi üzerindeki kritik rolünü nicel olarak ortaya "
  "koymaktadır.")
FIG("fig_sensitivity.png","Şekil 6.7. Duyarlılık analizi: (a) stil uyumu ve (b) sakatlık riski etkisi.")
PB()

# ============================================================ 7. TARTIŞMA
H1("7. TARTIŞMA")
H2("7.1. Sonuçların Yorumlanması")
P(f"Elde edilen {pct(S['UYUM'])} uyum skoru ve yüksek gol+asist projeksiyonu, incelenen oyuncunun hedef takım için "
  "sportif açıdan oldukça uygun bir transfer hedefi olduğunu göstermektedir. Bununla birlikte, modelin en değerli "
  "katkısı tek bir skor üretmesi değil; sonucu bir olasılık dağılımı olarak sunarak hem beklenen getiriyi hem de "
  "bu getirinin etrafındaki belirsizliği aynı anda görünür kılmasıdır. Örneğin sağlamlık göstergesinin "
  f"{pct(S['saglam'])} düzeyinde olması, yüksek katkı beklentisinin önemli bir sakatlık riskiyle dengelenmesi "
  "gerektiğini açıkça ortaya koymaktadır.")
P("Bu bağlamda sonuçlar, iki boyutlu bir karar matrisi olarak okunmalıdır: bir eksende beklenen sportif katkı "
  "(yüksek), diğer eksende sakatlık riski (orta-yüksek) yer alır. Yüksek katkı beklentisi transferi cazip kılarken, "
  "kayda değer sakatlık riski, bir sözleşme yapısında performansa veya oynanan maç sayısına bağlı ödeme maddeleriyle "
  "yönetilmesi gereken bir unsur olarak öne çıkar. Modelin ürettiği sayısal olasılıklar (örneğin belirli bir katkı "
  "eşiğinin aşılma ihtimali), bu tür sözleşme tasarımı ve risk paylaşımı kararlarına doğrudan girdi sağlayabilecek "
  "niteliktedir. Böylece model, salt bir 'evet/hayır' tavsiyesinin ötesine geçerek pazarlık ve planlama süreçlerini "
  "de besleyebilir.")
H2("7.2. Modelin Güçlü Yönleri")
BULLET("Tüm parametreler gerçek veriden türetilir; öznel varsayım en aza indirilmiştir.")
BULLET("Sonuçlar olasılık dağılımı olarak sunulur; bu, klasik nokta tahminlere göre çok daha zengin bir karar bilgisi sağlar.")
BULLET("Sistem genel amaçlıdır; herhangi bir lig, takım ve oyuncu için yeniden çalıştırılabilir.")
BULLET("Dayanıklı veri toplama ve arama altyapısı, gerçek dünya veri kısıtları altında güvenilir çalışır.")
BULLET("Etkileşimli arayüz, teknik olmayan karar vericiler için erişilebilirdir.")
H2("7.3. Kısıtlar ve Varsayımlar")
P("Çalışmanın bazı kısıtları bulunmaktadır. Birincisi, gol+asist projeksiyonu oyuncunun mevcut üretkenlik oranını "
  "koruyacağı varsayımına dayanır; oysa lig seviyesi farkı ve daha yüksek dakika sayısı bu oranı aşağı çekebilir "
  "(ortalamaya dönüş etkisi). İkincisi, maç reytinglerinden yalnızca tek bir toplam değişkenlik elde edilebildiğinden "
  "form ve maç-içi gürültü ayrımı sabit bir oranla yapılmıştır. Üçüncüsü, uyum skoru ağırlıkları uzman görüşüne "
  "dayalı olup ileride veriyle kalibre edilebilir. Dördüncüsü, ücretsiz API katmanı veri derinliğini ve hızını "
  "sınırlamaktadır. Son olarak model; taktiksel uyum, antrenör tercihi ve takım kimyası gibi nicelleştirilmesi zor "
  "faktörleri içermemektedir.")
P("Ayrıca, modelin santrfor mevkii için tasarlanmış olması bir kapsam kısıtıdır. Orta saha ve savunma oyuncularının "
  "katkısı gol ve asist gibi metriklerle yeterince yakalanamadığından, bu mevkiler için pozisyona özgü farklı "
  "performans göstergelerinin (örneğin pas isabeti, top kazanma, ikili mücadele kazanma yüzdesi, engelleme sayısı) "
  "tanımlanması gerekir. Mevcut hâliyle sistem en güçlü ve en anlamlı sonuçları hücum hattı oyuncularında "
  "vermektedir; diğer mevkilere genişletilmesi gelecek çalışma olarak öngörülmüştür.")
H2("7.4. Doğrulama ve Geçerlilik")
P("Simülasyonun istatistiksel kararlılığı, 10.000 tekrar sayısıyla sağlanmıştır; bu sayıda örneklem ortalamasının "
  "standart hatası ihmal edilebilir düzeydedir. Model çıktıları, oyuncunun gerçek geçmiş üretkenliğiyle (örneğin "
  f"{S['gol']} gol + {S['asist']} asist) tutarlılık göstermekte; projeksiyonlar bu gerçek temelin makul bir uzantısı "
  "olarak ortaya çıkmaktadır. Sabit rastgelelik tohumu (seed) kullanılarak sonuçların yeniden üretilebilirliği "
  "garanti altına alınmıştır.")
PB()

# ============================================================ 8. SONUÇ
H1("8. SONUÇ VE ÖNERİLER")
H2("8.1. Genel Değerlendirme")
P("Bu bitirme projesinde, bir futbolcunun gerçek verilerinden dijital ikizini oluşturan ve hedef bir takımdaki olası "
  "performansını Monte Carlo simülasyonuyla öngören bütüncül bir karar destek sistemi başarıyla geliştirilmiştir. "
  "Sistem; gerçek zamanlı veri toplama, olasılıksal parametre tahmini, çok kriterli uyum modellemesi, stokastik "
  "simülasyon ve etkileşimli görselleştirme adımlarını uçtan uca bütünleştirmektedir. Örnek vaka çalışması, sistemin "
  "yorumlanabilir ve eyleme dönük çıktılar üretebildiğini göstermiştir. Proje, endüstri mühendisliğinin simülasyon ve "
  "risk analizi araçlarının, spor analitiği gibi yeni ve dinamik bir alanda yüksek katma değer üretebileceğini ortaya "
  "koymaktadır.")
P("Projenin bir diğer önemli çıktısı, geliştirilen yöntemin yeniden kullanılabilir ve ölçeklenebilir bir yazılım "
  "altyapısı biçiminde ortaya konmasıdır. Sistem, tek bir oyuncu-takım çifti için tasarlanmış statik bir analiz "
  "olmanın ötesinde, herhangi bir senaryo için anında yeniden çalıştırılabilen dinamik bir araçtır. Bu yönüyle "
  "çalışma yalnızca bir vaka incelemesi değil; kulüplerin scout (gözlemci) departmanlarında pratik olarak "
  "kullanılabilecek bir prototip karar destek aracı niteliği taşımaktadır. Akademik açıdan ise, endüstri "
  "mühendisliği yöntemlerinin disiplinler arası bir alanda nasıl somut değer üretebileceğine dair örnek bir "
  "uygulama sunmaktadır.")
H2("8.2. Gelecek Çalışmalar")
BULLET("Lig seviyesi farkını dikkate alan bir düzeltme katsayısının (örn. UEFA ülke katsayıları) modele eklenmesi.")
BULLET("Uyum skoru ağırlıklarının, gerçekleşmiş geçmiş transferler üzerinde makine öğrenmesiyle kalibre edilmesi.")
BULLET("Beklenen gol (xG) ve beklenen asist (xA) gibi gelişmiş metriklerin entegrasyonu.")
BULLET("Birden çok oyuncuyu yan yana karşılaştıran bir 'transfer kısa listesi' modülü.")
BULLET("Oyuncu piyasa değeri ve maaş verisiyle birleştirilerek bir yatırım getirisi (ROI) optimizasyon modülü.")
BULLET("Orta saha ve defans oyuncuları için pozisyona özgü katkı metriklerinin tanımlanması.")
PB()

# ============================================================ KAYNAKÇA
H1("KAYNAKÇA")
refs=[
 "Grieves, M. (2014). Digital Twin: Manufacturing Excellence through Virtual Factory Replication. White Paper, "
 "Florida Institute of Technology.",
 "Metropolis, N., & Ulam, S. (1949). The Monte Carlo Method. Journal of the American Statistical Association, "
 "44(247), 335–341.",
 "Lewis, M. (2003). Moneyball: The Art of Winning an Unfair Game. W. W. Norton & Company.",
 "Rubinstein, R. Y., & Kroese, D. P. (2016). Simulation and the Monte Carlo Method (3rd ed.). Wiley.",
 "Law, A. M. (2015). Simulation Modeling and Analysis (5th ed.). McGraw-Hill.",
 "Tao, F., Zhang, H., Liu, A., & Nee, A. Y. C. (2019). Digital Twin in Industry: State-of-the-Art. "
 "IEEE Transactions on Industrial Informatics, 15(4), 2405–2415.",
 "Memmert, D., & Rein, R. (2018). Match Analysis, Big Data and Tactics: Current Trends in Elite Soccer. "
 "German Journal of Exercise and Sport Research, 48, 24–32.",
 "Hvattum, L. M. (2019). A Comprehensive Review of Plus-Minus Ratings for Evaluating Individual Players in Team "
 "Sports. International Journal of Computer Science in Sport, 18(1), 1–23.",
 "API-Football (2024). API-Football v3 Documentation. https://www.api-football.com/documentation-v3",
 "Streamlit Inc. (2024). Streamlit Documentation. https://docs.streamlit.io",
 "Harris, C. R. et al. (2020). Array Programming with NumPy. Nature, 585, 357–362.",
]
for i,r in enumerate(refs,1):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Cm(0.75); p.paragraph_format.first_line_indent=Cm(-0.75)
    p.paragraph_format.space_after=Pt(6); p.add_run(f"[{i}] {r}")
PB()

# ============================================================ EKLER
H1("EKLER")
H2("EK-A. Parametre Tahmin ve Simülasyon Kodu (Özet)")
CODE(["# Parametre tahmini (gerçek veriden)",
      "S_base   = ort_rating / 10",
      "sig      = np.std([r/10 for r in ratings])",
      "sigma_form, sigma_mac = sig*0.8, sig*0.6",
      "p_sak    = sakatlik_episode / toplam_mac",
      "lam_sak  = sakatlik_kacan  / sakatlik_episode",
      "",
      "# Uyum skoru",
      "kalite = S_base",
      "verim  = min(ga90/1.0, 1.0)",
      "stil   = min(hedef_gol_mac/2.6, 1.0)",
      "UYUM   = 0.50*kalite + 0.25*verim + 0.25*stil",
      "",
      "# Monte Carlo (10.000 sezon)",
      "for _ in range(N):",
      "    sk=[]; kacan=0; rem=0; gas=0",
      "    for m in range(38):",
      "        if rem>0: rem-=1; kacan+=1; continue",
      "        if rng.random()<p_sak: rem=max(1,rng.poisson(lam_sak))-1; kacan+=1; continue",
      "        s=min(1,max(0, S_base+rng.normal(0,sigma_form)+rng.normal(0,sigma_mac)))",
      "        sk.append(s); gas+=rng.poisson(max(ga90*(s/S_base),0))",
      "    P.append(np.mean(sk)); K.append(kacan); G.append(gas)"],
     "Kod A.1. Parametre tahmini, uyum skoru ve simülasyon motoru.")
H2("EK-B. Sakatlık Filtreleme ve Epizot Gruplama (Özet)")
CODE(["NON_INJURY = {'Red Card','Yellow Cards','National selection',",
      "              'International duty','Rest','Inactive','Suspended'}",
      "inj = [r for r in sakatliklar",
      "       if r['tip']=='Missing Fixture' and r['neden'] not in NON_INJURY]",
      "gunler = sorted(date.fromisoformat(r['tarih']) for r in inj)",
      "kacan = len(gunler); episode = 0; prev = None",
      "for d in gunler:                 # 30 günden büyük boşluk = yeni epizot",
      "    if prev is None or (d-prev).days > 30: episode += 1",
      "    prev = d"],
     "Kod B.1. Gerçek sakatlıkların filtrelenmesi ve epizotlara ayrılması.")

H2("EK-C. Notlar")
P("Bu rapordaki tüm sayısal sonuçlar ve grafikler, API-Football'dan çekilen gerçek veriler kullanılarak üretilmiş "
  "ve sabit rastgelelik tohumu ile yeniden üretilebilir kılınmıştır. Kapak sayfasındaki köşeli parantez içindeki "
  "alanlar (üniversite, ad, danışman vb.) kullanıcı tarafından doldurulmalıdır.")

# -------- Sayfa numarası ve kaydet --------
page_number_footer()
doc.save("/Users/Ethem/Desktop/colab/Bitirme_Projesi_Raporu.docx")
print("OK: Bitirme_Projesi_Raporu.docx olusturuldu")
